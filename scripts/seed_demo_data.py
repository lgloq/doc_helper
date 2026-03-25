from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parents[1]
BACKEND_DIR = ROOT_DIR / "backend"
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from docx import Document as DocxDocument
from fastapi import UploadFile
from sqlalchemy import select

from app.db.session import SessionLocal
from app.models.document import Document, DocumentVersion
from app.models.enums import DocumentStatus, IngestStatus, RoleName
from app.models.user import User
from app.repositories.user_repository import UserRepository
from app.schemas.document import DocumentACLCreate, DocumentIngestRequest
from app.services.auth.bootstrap import seed_mock_data
from app.services.documents.service import DocumentService
from app.services.eval.bootstrap import seed_demo_eval_cases
from app.services.ingestion.service import DocumentIngestionService


def write_demo_assets(base_dir: Path) -> dict[str, Path]:
    assets: dict[str, Path] = {}

    public_handbook = base_dir / "public_handbook.md"
    public_handbook.write_text(
        "# Company Handbook\n\n"
        "## Holiday Schedule\n"
        "All employees follow the official holiday schedule published by People Ops.\n"
        "Requests for exceptions should be submitted at least one week in advance.\n\n"
        "## Support Window\n"
        "Managers should communicate any holiday coverage changes before the end of the week.\n",
        encoding="utf-8",
    )
    assets["public_handbook"] = public_handbook

    platform_runbook_v1 = base_dir / "platform_runbook_v1.txt"
    platform_runbook_v1.write_text(
        "Platform release checklist\n\n"
        "1. Confirm rollout owner and change window.\n"
        "2. Validate deployment artifacts in staging.\n"
        "3. Notify stakeholders before production rollout.\n"
        "4. Record rollback contact list in the release ticket.\n",
        encoding="utf-8",
    )
    assets["platform_runbook_v1"] = platform_runbook_v1

    platform_runbook_v2 = base_dir / "platform_runbook_v2.txt"
    platform_runbook_v2.write_text(
        "Platform release checklist\n\n"
        "1. Confirm rollout owner, change window, and incident commander.\n"
        "2. Validate deployment artifacts in staging and confirm monitoring alerts.\n"
        "3. Notify stakeholders before production rollout and publish rollback steps.\n"
        "4. Record rollback contact list and verification checklist in the release ticket.\n"
        "5. If rollback exceeds 15 minutes, escalate to the platform manager immediately.\n",
        encoding="utf-8",
    )
    assets["platform_runbook_v2"] = platform_runbook_v2

    incident_guide = base_dir / "incident_response_guide.docx"
    doc = DocxDocument()
    doc.add_heading("Incident Response Guide", level=1)
    doc.add_paragraph("Managers should acknowledge customer-facing incidents within fifteen minutes.")
    doc.add_paragraph("The owner must update the incident timeline and assign a communications lead.")
    doc.save(incident_guide)
    assets["incident_guide"] = incident_guide

    security_exceptions = base_dir / "security_exceptions.html"
    security_exceptions.write_text(
        "<html><body>"
        "<h1>Security Exceptions</h1>"
        "<p>This document is restricted to administrators and contains temporary access exception notes.</p>"
        "<p>Do not share exception tokens outside the security review channel.</p>"
        "</body></html>",
        encoding="utf-8",
    )
    assets["security_exceptions"] = security_exceptions

    return assets


def find_document_by_title(session, title: str) -> Document | None:
    statement = select(Document).where(Document.title == title).order_by(Document.created_at.desc())
    return session.scalar(statement)


def list_versions(session, document_id) -> list[DocumentVersion]:
    statement = (
        select(DocumentVersion)
        .where(DocumentVersion.document_id == document_id)
        .order_by(DocumentVersion.version_number.asc())
    )
    return list(session.scalars(statement).all())


def upload_document_from_path(
    ingestion_service: DocumentIngestionService,
    actor: User,
    path: Path,
    *,
    title: str,
    description: str,
):
    with path.open("rb") as handle:
        upload = UploadFile(file=handle, filename=path.name)
        response = ingestion_service.upload_document(actor, upload, title, description, DocumentStatus.ACTIVE)
    ingestion_service.ingest_document(actor, response.document.id, DocumentIngestRequest(version_id=response.version.id))
    return response.document.id


def upload_document_version_from_path(
    ingestion_service: DocumentIngestionService,
    actor: User,
    document_id,
    path: Path,
) -> None:
    with path.open("rb") as handle:
        upload = UploadFile(file=handle, filename=path.name)
        response = ingestion_service.upload_document_version(actor, document_id, upload)
    ingestion_service.ingest_document(actor, document_id, DocumentIngestRequest(version_id=response.version.id))


def ensure_acl(document_service: DocumentService, actor: User, document_id, payload: DocumentACLCreate) -> None:
    document_service.upsert_acl_entry(actor, document_id, payload)


def ensure_ready_current_version(ingestion_service: DocumentIngestionService, actor: User, document_id, version_id) -> None:
    ingestion_service.ingest_document(actor, document_id, DocumentIngestRequest(version_id=version_id))


def main() -> None:
    seed_mock_data()
    seed_demo_eval_cases()

    with tempfile.TemporaryDirectory(prefix="eka-demo-assets-") as temp_dir:
        assets = write_demo_assets(Path(temp_dir))
        session = SessionLocal()
        try:
            user_repository = UserRepository(session)
            admin_user = user_repository.get_by_email("admin@local.test")
            if admin_user is None:
                raise SystemExit("admin@local.test was not found. Seed users first.")

            ingestion_service = DocumentIngestionService(session)
            document_service = DocumentService(session)

            public_document = find_document_by_title(session, "Public Handbook")
            if public_document is None:
                public_document_id = upload_document_from_path(
                    ingestion_service,
                    admin_user,
                    assets["public_handbook"],
                    title="Public Handbook",
                    description="Public-facing handbook content for all employees.",
                )
            else:
                public_document_id = public_document.id
                public_versions = list_versions(session, public_document_id)
                if public_versions and public_versions[-1].ingest_status != IngestStatus.READY:
                    ensure_ready_current_version(ingestion_service, admin_user, public_document_id, public_versions[-1].id)
            ensure_acl(
                document_service,
                admin_user,
                public_document_id,
                DocumentACLCreate(principal_type="public", can_view=True, can_manage=False),
            )

            platform_document = find_document_by_title(session, "Platform Runbook")
            if platform_document is None:
                platform_document_id = upload_document_from_path(
                    ingestion_service,
                    admin_user,
                    assets["platform_runbook_v1"],
                    title="Platform Runbook",
                    description="Team-scoped release checklist and deployment guidance.",
                )
            else:
                platform_document_id = platform_document.id
                platform_versions = list_versions(session, platform_document_id)
                if platform_versions and platform_versions[-1].ingest_status != IngestStatus.READY:
                    ensure_ready_current_version(ingestion_service, admin_user, platform_document_id, platform_versions[-1].id)
            ensure_acl(
                document_service,
                admin_user,
                platform_document_id,
                DocumentACLCreate(principal_type="team", team_name="platform", can_view=True, can_manage=False),
            )
            if len(list_versions(session, platform_document_id)) < 2:
                upload_document_version_from_path(ingestion_service, admin_user, platform_document_id, assets["platform_runbook_v2"])

            incident_document = find_document_by_title(session, "Incident Response Guide")
            if incident_document is None:
                incident_document_id = upload_document_from_path(
                    ingestion_service,
                    admin_user,
                    assets["incident_guide"],
                    title="Incident Response Guide",
                    description="Manager guidance for handling customer-facing incidents.",
                )
            else:
                incident_document_id = incident_document.id
                incident_versions = list_versions(session, incident_document_id)
                if incident_versions and incident_versions[-1].ingest_status != IngestStatus.READY:
                    ensure_ready_current_version(ingestion_service, admin_user, incident_document_id, incident_versions[-1].id)
            ensure_acl(
                document_service,
                admin_user,
                incident_document_id,
                DocumentACLCreate(principal_type="role", role_name=RoleName.MANAGER, can_view=True, can_manage=False),
            )

            security_document = find_document_by_title(session, "Security Exceptions")
            if security_document is None:
                security_document_id = upload_document_from_path(
                    ingestion_service,
                    admin_user,
                    assets["security_exceptions"],
                    title="Security Exceptions",
                    description="Admin-only exception register used to demonstrate isolation.",
                )
            else:
                security_document_id = security_document.id
                security_versions = list_versions(session, security_document_id)
                if security_versions and security_versions[-1].ingest_status != IngestStatus.READY:
                    ensure_ready_current_version(ingestion_service, admin_user, security_document_id, security_versions[-1].id)

            session.commit()

            print("Demo data is ready.")
            print("Seeded documents:")
            print(f"- Public Handbook (public ACL) -> {public_document_id}")
            print(f"- Platform Runbook (team ACL: platform, with 2 versions) -> {platform_document_id}")
            print(f"- Incident Response Guide (manager role ACL) -> {incident_document_id}")
            print(f"- Security Exceptions (owner/admin only) -> {security_document_id}")
            print("")
            print("Suggested demo accounts:")
            print("- viewer@local.test / viewer123")
            print("- manager@local.test / manager123")
            print("- admin@local.test / admin123")
            print("")
            print("Suggested demo questions:")
            print("- What does the platform release checklist require?")
            print("- What does the company handbook say about holiday schedule?")
            print("- What should managers do during a customer-facing incident?")
        finally:
            session.close()


if __name__ == "__main__":
    main()
