from __future__ import annotations

import base64
import csv
import re
from dataclasses import dataclass, replace
from pathlib import Path
from urllib.parse import unquote, urlparse

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.core.config import get_settings
from app.services.ingestion.ocr import OcrResult, OcrService
from app.services.ingestion.table_utils import table_rows_to_text_segments

try:
    import pdfplumber
except ImportError:  # pragma: no cover - optional until the backend image is rebuilt
    pdfplumber = None

try:
    import fitz
except ImportError:  # pragma: no cover - optional until OCR dependencies are installed
    fitz = None


IMAGE_FILE_SUFFIXES = {".png", ".jpg", ".jpeg"}
HTML_CONTENT_SELECTORS = (
    ".UCAP-CONTENT",
    ".pages_content",
    ".TRS_UEDITOR",
    ".trs_editor_view",
    ".BodyLabel",
    ".main-content",
    ".policyLibraryOverview_content",
    ".my_doccontent",
    ".TRS_Editor",
    "main",
    "article",
    "[role=main]",
    ".article",
    ".content",
    "#content",
)
HTML_NOISE_LINE_PATTERNS = (
    r"^(首页|登录|注册|退出|个人中心|邮箱|无障碍|EN|简|繁|默认|大|超大|关闭|返回|回到顶部)$",
    r"^(打印|收藏|留言|分享|分享到微信朋友圈|扫一扫|纠错|客户端|微博|微信|二维码|热门检索)[:：]?.*$",
    r"^当前位置[:：]?.*$",
    r"^您现在的位置[:：]?.*$",
    r"^字号[:：]?.*$",
    r"^【\s*(大|中|小|打印此页|关闭窗口)\s*】$",
    r"^网站标识码.*$",
    r"^京ICP备.*$",
    r"^京公网安备.*$",
    r"^版权所有.*$",
)
HTML_NOISE_TERMS = (
    "登录",
    "注册",
    "热门检索",
    "分享",
    "扫一扫",
    "打印",
    "首页",
    "当前位置",
    "个人中心",
    "无障碍",
    "客户端",
    "网站标识码",
)


@dataclass
class ParsedSegment:
    text: str
    page_number: int | None = None
    paragraph_index: int | None = None
    section_title: str | None = None
    char_start: int | None = None
    char_end: int | None = None


@dataclass
class ParsedDocument:
    normalized_text: str
    segments: list[ParsedSegment]
    page_count: int | None
    parser_name: str


@dataclass
class PdfPlumberExtraction:
    page_count: int
    page_texts: dict[int, str]
    table_segments_by_page: dict[int, list[tuple[str, str]]]


class DocumentParser:
    def __init__(self, ocr_service: OcrService | None = None):
        self.settings = get_settings()
        self.ocr_service = ocr_service or OcrService(self.settings)

    def parse(self, path: Path) -> ParsedDocument:
        suffix = path.suffix.lower()
        if suffix == ".txt":
            return self._parse_txt(path)
        if suffix in {".md", ".markdown"}:
            return self._parse_markdown(path)
        if suffix in {".html", ".htm"}:
            return self._parse_html(path)
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix == ".csv":
            return self._parse_csv(path)
        if suffix in IMAGE_FILE_SUFFIXES:
            return self._parse_image(path)
        raise ValueError(f"Unsupported parser for file type '{suffix}'.")

    def _parse_txt(self, path: Path) -> ParsedDocument:
        text = self._read_text_file(path)
        segments = []
        for idx, paragraph in enumerate(self._split_plain_paragraphs(text), start=1):
            segments.append(ParsedSegment(text=paragraph, paragraph_index=idx))
        return self._finalize_segments(segments, parser_name="txt")

    def _parse_markdown(self, path: Path) -> ParsedDocument:
        raw_text = self._read_text_file(path)
        lines = raw_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        segments: list[ParsedSegment] = []
        current_section: str | None = None
        paragraph_buffer: list[str] = []
        paragraph_index = 0

        def flush_buffer() -> None:
            nonlocal paragraph_index, paragraph_buffer
            if not paragraph_buffer:
                return
            paragraph_index += 1
            paragraph = self._normalize_text(" ".join(paragraph_buffer))
            if paragraph:
                segments.append(
                    ParsedSegment(
                        text=paragraph,
                        paragraph_index=paragraph_index,
                        section_title=current_section,
                    )
                )
            paragraph_buffer = []

        index = 0
        while index < len(lines):
            line = lines[index]
            stripped = line.strip()
            heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading_match:
                flush_buffer()
                current_section = self._strip_markdown_inline(heading_match.group(2))
                paragraph_index += 1
                segments.append(
                    ParsedSegment(
                        text=current_section,
                        paragraph_index=paragraph_index,
                        section_title=current_section,
                    )
                )
                index += 1
                continue
            if self._is_markdown_table_start(lines, index):
                flush_buffer()
                table_rows, next_index = self._collect_markdown_table(lines, index)
                for text in table_rows_to_text_segments(table_rows, caption=current_section):
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=text,
                            paragraph_index=paragraph_index,
                            section_title=current_section,
                        )
                    )
                index = next_index
                continue
            image_sources = self._extract_markdown_image_sources(stripped)
            if image_sources:
                image_line_text = self._strip_markdown_inline(stripped)
                if image_line_text:
                    paragraph_buffer.append(image_line_text)
                flush_buffer()
                paragraph_index = self._append_image_source_segments(
                    segments,
                    image_sources,
                    paragraph_index=paragraph_index,
                    source_path=path,
                    section_title=current_section,
                    section_title_prefix="Markdown image",
                )
                index += 1
                continue
            if not stripped:
                flush_buffer()
                index += 1
                continue
            paragraph_buffer.append(self._strip_markdown_inline(stripped))
            index += 1

        flush_buffer()
        return self._finalize_segments(segments, parser_name="markdown")

    def _parse_html(self, path: Path) -> ParsedDocument:
        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
        for element in soup(["script", "style", "noscript", "template", "svg"]):
            element.decompose()
        segments: list[ParsedSegment] = []
        current_section: str | None = None
        paragraph_index = 0

        content_root = self._select_html_content_root(soup)
        for element in content_root.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table", "img"]):
            if self._is_html_boilerplate_element(element):
                continue
            if element.name == "img":
                src = element.get("src")
                if src:
                    paragraph_index = self._append_image_source_segments(
                        segments,
                        [src],
                        paragraph_index=paragraph_index,
                        source_path=path,
                        section_title=current_section,
                        section_title_prefix="HTML image",
                    )
                continue
            if element.name == "table":
                table_visible_text = self._normalize_html_text(element.get_text("\n", strip=True))
                if self._is_html_metadata_table(table_visible_text):
                    continue
                table_rows = self._extract_html_table_rows(element)
                caption = self._normalize_text(element.find("caption").get_text(" ", strip=True)) if element.find("caption") else current_section
                for text in table_rows_to_text_segments(table_rows, caption=caption):
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=text,
                            paragraph_index=paragraph_index,
                            section_title=current_section,
                        )
                    )
                continue
            text = self._normalize_html_text(element.get_text("\n", strip=True))
            if not text:
                continue
            if element.name and element.name.startswith("h"):
                current_section = text
            paragraph_index += 1
            segments.append(
                ParsedSegment(
                    text=text,
                    paragraph_index=paragraph_index,
                    section_title=current_section,
                )
            )

        if not segments:
            fallback_text = self._normalize_text(soup.get_text("\n", strip=True))
            segments = [ParsedSegment(text=fallback_text, paragraph_index=1)] if fallback_text else []

        return self._finalize_segments(segments, parser_name="html")

    def _select_html_content_root(self, soup: BeautifulSoup):
        candidates = []
        seen_ids: set[int] = set()
        for selector in HTML_CONTENT_SELECTORS:
            for root in soup.select(selector):
                root_id = id(root)
                if root_id in seen_ids:
                    continue
                seen_ids.add(root_id)
                score = self._score_html_content_root(root)
                if score > 0:
                    candidates.append((score, root))

        if candidates:
            candidates.sort(key=lambda item: item[0], reverse=True)
            return candidates[0][1]
        return soup.body or soup

    @classmethod
    def _score_html_content_root(cls, root) -> float:
        text = cls._normalize_html_text(root.get_text("\n", strip=True))
        if not text:
            return 0.0
        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
        if chinese_chars < 40:
            return 0.0
        link_text_chars = sum(len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", link.get_text(" ", strip=True))) for link in root.find_all("a"))
        noise_hits = sum(text.count(term) for term in HTML_NOISE_TERMS)
        paragraph_like_lines = sum(1 for line in text.split("\n") if len(cls._normalize_text(line)) >= 20)
        return chinese_chars + paragraph_like_lines * 40 - link_text_chars * 2.0 - noise_hits * 120.0

    @staticmethod
    def _is_html_boilerplate_element(element) -> bool:
        for parent in [element, *element.parents]:
            name = getattr(parent, "name", None)
            if name in {"nav", "header", "footer", "aside"}:
                return True
            attrs = []
            for key in ("class", "id", "role", "aria-label"):
                value = parent.get(key) if hasattr(parent, "get") else None
                if isinstance(value, list):
                    attrs.extend(str(item).casefold() for item in value)
                elif value:
                    attrs.append(str(value).casefold())
            attr_text = " ".join(attrs)
            if re.search(r"\b(nav|navbar|breadcrumb|footer|header|sidebar|menu|toolbar|pagination|copyright)\b", attr_text):
                return True
        return False

    @classmethod
    def _normalize_html_text(cls, text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        kept_lines: list[str] = []
        seen_lines: set[str] = set()
        for raw_line in normalized.split("\n"):
            line = cls._normalize_text(raw_line)
            if not line or cls._is_html_noise_line(line):
                continue
            dedupe_key = cls._normalize_matchable_html_line(line)
            if dedupe_key in seen_lines:
                continue
            seen_lines.add(dedupe_key)
            kept_lines.append(line)
        return "\n".join(kept_lines)

    @staticmethod
    def _is_html_noise_line(line: str) -> bool:
        compact = re.sub(r"\s+", "", line)
        if not compact:
            return True
        if len(compact) <= 2 and not re.search(r"[\u4e00-\u9fff]{2,}", compact):
            return True
        if any(re.fullmatch(pattern, compact, flags=re.IGNORECASE) for pattern in HTML_NOISE_LINE_PATTERNS):
            return True
        if len(compact) <= 12 and any(term in compact for term in HTML_NOISE_TERMS):
            return True
        return False

    @staticmethod
    def _is_html_metadata_table(text: str) -> bool:
        if not text:
            return False
        metadata_labels = (
            "索引号",
            "主题分类",
            "发文机关",
            "成文日期",
            "标题",
            "发文字号",
            "发布日期",
            "来源",
            "来 源",
        )
        hits = sum(1 for label in metadata_labels if label in text.replace(" ", ""))
        return hits >= 3 and len(text) < 1200

    @staticmethod
    def _normalize_matchable_html_line(line: str) -> str:
        return re.sub(r"[\s|｜]+", "", line).casefold()

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        pdfplumber_extraction = self._extract_pdf_with_pdfplumber(path)
        if pdfplumber_extraction and (
            self.ocr_service.enabled
            or any(text.strip() for text in pdfplumber_extraction.page_texts.values())
            or any(pdfplumber_extraction.table_segments_by_page.values())
        ):
            return self._parse_pdf_pdfplumber_pages(path, pdfplumber_extraction)

        reader = PdfReader(str(path))
        segments: list[ParsedSegment] = []
        paragraph_index = 0
        pdf_image_document = self._open_pdf_image_document(path)
        try:
            for page_number, page in enumerate(reader.pages, start=1):
                raw_text = page.extract_text() or ""
                paragraphs = self._split_plain_paragraphs(raw_text)
                if not paragraphs and raw_text.strip():
                    paragraphs = [self._normalize_text(raw_text)]
                for paragraph in paragraphs:
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=paragraph,
                            page_number=page_number,
                            paragraph_index=paragraph_index,
                        )
                    )
                page_segment_count = len(paragraphs)
                should_ocr_page = self._should_ocr_pdf_page(
                    page_number=page_number,
                    raw_text=raw_text,
                    page_segment_count=page_segment_count,
                    has_table_segments=False,
                )
                if should_ocr_page:
                    paragraph_index = self._append_ocr_segments(
                        segments,
                        self.ocr_service.extract_pdf_page(path, page_number),
                        paragraph_index=paragraph_index,
                        page_number=page_number,
                        ocr_section_title=f"PDF page {page_number} OCR",
                        table_caption_prefix=f"PDF page {page_number} OCR table",
                    )
                    continue
                paragraph_index = self._append_pdf_embedded_image_segments(
                    segments,
                    pdf_image_document,
                    paragraph_index=paragraph_index,
                    page_number=page_number,
                )
        finally:
            if pdf_image_document is not None and hasattr(pdf_image_document, "close"):
                pdf_image_document.close()

        return self._finalize_segments(segments, parser_name="pdf", page_count=len(reader.pages))

    def _parse_pdf_pdfplumber_pages(self, path: Path, extraction: PdfPlumberExtraction) -> ParsedDocument:
        segments: list[ParsedSegment] = []
        paragraph_index = 0
        pdf_image_document = self._open_pdf_image_document(path)
        try:
            for page_number in range(1, extraction.page_count + 1):
                page_segment_start = len(segments)
                raw_text = extraction.page_texts.get(page_number, "")
                paragraphs = self._split_plain_paragraphs(raw_text)
                if not paragraphs and raw_text.strip():
                    paragraphs = [self._normalize_text(raw_text)]
                for paragraph in paragraphs:
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=paragraph,
                            page_number=page_number,
                            paragraph_index=paragraph_index,
                        )
                    )
                table_segments = extraction.table_segments_by_page.get(page_number, [])
                for table_text, section_title in table_segments:
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=table_text,
                            page_number=page_number,
                            paragraph_index=paragraph_index,
                            section_title=section_title,
                        )
                    )
                should_ocr_page = self._should_ocr_pdf_page(
                    page_number=page_number,
                    raw_text=raw_text,
                    page_segment_count=len(segments) - page_segment_start,
                    has_table_segments=bool(table_segments),
                )
                if should_ocr_page:
                    paragraph_index = self._append_ocr_segments(
                        segments,
                        self.ocr_service.extract_pdf_page(path, page_number),
                        paragraph_index=paragraph_index,
                        page_number=page_number,
                        ocr_section_title=f"PDF page {page_number} OCR",
                        table_caption_prefix=f"PDF page {page_number} OCR table",
                    )
                    continue
                paragraph_index = self._append_pdf_embedded_image_segments(
                    segments,
                    pdf_image_document,
                    paragraph_index=paragraph_index,
                    page_number=page_number,
                )
        finally:
            if pdf_image_document is not None and hasattr(pdf_image_document, "close"):
                pdf_image_document.close()
        return self._finalize_segments(segments, parser_name="pdf", page_count=extraction.page_count)

    def _parse_docx(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(str(path))
        segments: list[ParsedSegment] = []
        paragraph_index = 0
        current_section: str | None = None

        for block in self._iter_docx_blocks(doc):
            if isinstance(block, Paragraph):
                paragraph_images = self._extract_docx_paragraph_image_payloads(doc, block)
                text = self._normalize_text(block.text)
                if text:
                    style_name = (block.style.name or "").lower() if block.style else ""
                    if style_name.startswith("heading"):
                        current_section = text
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=text,
                            paragraph_index=paragraph_index,
                            section_title=current_section,
                        )
                    )
                if paragraph_images:
                    paragraph_index = self._append_image_payload_segments(
                        segments,
                        paragraph_images,
                        paragraph_index=paragraph_index,
                        section_title=current_section,
                        section_title_prefix="DOCX image",
                    )
                continue
            if isinstance(block, Table):
                table_rows = [
                    [self._normalize_text(cell.text) for cell in row.cells]
                    for row in block.rows
                ]
                for text in table_rows_to_text_segments(table_rows, caption=current_section):
                    paragraph_index += 1
                    segments.append(
                        ParsedSegment(
                            text=text,
                            paragraph_index=paragraph_index,
                            section_title=current_section,
                        )
                    )
                table_images = self._extract_docx_table_image_payloads(doc, block)
                if table_images:
                    paragraph_index = self._append_image_payload_segments(
                        segments,
                        table_images,
                        paragraph_index=paragraph_index,
                        section_title=current_section,
                        section_title_prefix="DOCX image",
                    )

        return self._finalize_segments(segments, parser_name="docx")

    def _parse_csv(self, path: Path) -> ParsedDocument:
        text = self._read_text_file(path)
        rows = list(csv.reader(text.splitlines()))
        segments = [
            ParsedSegment(text=segment, paragraph_index=index)
            for index, segment in enumerate(table_rows_to_text_segments(rows, caption=path.stem), start=1)
        ]
        return self._finalize_segments(segments, parser_name="csv")

    def _parse_image(self, path: Path) -> ParsedDocument:
        segments: list[ParsedSegment] = []
        ocr_result = self._filter_image_ocr_result(self.ocr_service.extract_image(path))
        self._append_ocr_segments(
            segments,
            ocr_result,
            paragraph_index=0,
            ocr_section_title="Image OCR",
            table_caption_prefix="Image table",
        )
        return self._finalize_segments(segments, parser_name="image", page_count=1)

    def _should_ocr_pdf_page(
        self,
        *,
        page_number: int,
        raw_text: str,
        page_segment_count: int,
        has_table_segments: bool,
    ) -> bool:
        if not self.ocr_service.enabled:
            return False
        if self.settings.ocr_max_pages > 0 and page_number > self.settings.ocr_max_pages:
            return False
        if page_segment_count == 0:
            return True
        if has_table_segments:
            return False
        return len(self._normalize_text(raw_text)) < self.settings.ocr_min_text_chars

    def _append_ocr_segments(
        self,
        segments: list[ParsedSegment],
        ocr_result: OcrResult,
        *,
        paragraph_index: int,
        page_number: int | None = None,
        ocr_section_title: str,
        table_caption_prefix: str,
    ) -> int:
        if ocr_result.text:
            paragraph_index += 1
            segments.append(
                ParsedSegment(
                    text=ocr_result.text,
                    page_number=page_number,
                    paragraph_index=paragraph_index,
                    section_title=ocr_section_title,
                )
            )

        for table_index, rows in enumerate(ocr_result.tables, start=1):
            section_title = f"{table_caption_prefix} {table_index}"
            for text in table_rows_to_text_segments(rows, caption=section_title):
                paragraph_index += 1
                segments.append(
                    ParsedSegment(
                        text=text,
                        page_number=page_number,
                        paragraph_index=paragraph_index,
                        section_title=section_title,
                    )
                )
        return paragraph_index

    def _append_image_source_segments(
        self,
        segments: list[ParsedSegment],
        image_sources: list[str],
        *,
        paragraph_index: int,
        source_path: Path,
        section_title: str | None,
        section_title_prefix: str,
        page_number: int | None = None,
    ) -> int:
        payloads = [payload for source in image_sources if (payload := self._resolve_image_payload(source_path, source)) is not None]
        return self._append_image_payload_segments(
            segments,
            payloads,
            paragraph_index=paragraph_index,
            section_title=section_title,
            section_title_prefix=section_title_prefix,
            page_number=page_number,
        )

    def _append_image_payload_segments(
        self,
        segments: list[ParsedSegment],
        image_payloads: list[bytes],
        *,
        paragraph_index: int,
        section_title: str | None,
        section_title_prefix: str,
        page_number: int | None = None,
    ) -> int:
        image_index = self._next_embedded_image_index(segments, section_title_prefix)
        for payload in image_payloads:
            ocr_result = self._filter_image_ocr_result(self.ocr_service.extract_bytes(payload))
            if not self._ocr_result_has_meaningful_content(ocr_result):
                continue
            base_title = f"{section_title_prefix} {image_index}"
            paragraph_index = self._append_ocr_segments(
                segments,
                ocr_result,
                paragraph_index=paragraph_index,
                page_number=page_number,
                ocr_section_title=f"{base_title} OCR",
                table_caption_prefix=f"{base_title} table",
            )
            image_index += 1
        return paragraph_index

    def _append_pdf_embedded_image_segments(
        self,
        segments: list[ParsedSegment],
        pdf_document,
        *,
        paragraph_index: int,
        page_number: int,
    ) -> int:
        if pdf_document is None:
            return paragraph_index
        payloads = self._extract_pdf_page_image_payloads(pdf_document, page_number)
        return self._append_image_payload_segments(
            segments,
            payloads,
            paragraph_index=paragraph_index,
            section_title=None,
            section_title_prefix=f"PDF page {page_number} image",
            page_number=page_number,
        )

    def _open_pdf_image_document(self, path: Path):
        if not self.ocr_service.enabled or fitz is None:
            return None
        try:
            return fitz.open(str(path))
        except Exception:
            return None

    def _extract_pdf_page_image_payloads(self, pdf_document, page_number: int) -> list[bytes]:
        if pdf_document is None:
            return []
        try:
            page = pdf_document.load_page(page_number - 1)
        except Exception:
            return []

        payloads: list[bytes] = []
        seen_xrefs: set[int] = set()
        for image_info in page.get_images(full=True):
            xref = image_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                image_info_payload = pdf_document.extract_image(xref)
                payload = image_info_payload.get("image")
            except Exception:
                continue
            if payload:
                payloads.append(payload)
        return payloads

    @staticmethod
    def _iter_docx_blocks(doc):
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc)

    def _extract_docx_paragraph_image_payloads(self, doc, paragraph: Paragraph) -> list[bytes]:
        return self._extract_docx_blip_payloads(doc, paragraph._element)

    def _extract_docx_table_image_payloads(self, doc, table: Table) -> list[bytes]:
        payloads: list[bytes] = []
        seen_payloads: set[bytes] = set()
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for payload in self._extract_docx_paragraph_image_payloads(doc, paragraph):
                        if payload in seen_payloads:
                            continue
                        seen_payloads.add(payload)
                        payloads.append(payload)
        return payloads

    @staticmethod
    def _extract_docx_blip_payloads(doc, element) -> list[bytes]:
        payloads: list[bytes] = []
        seen_rids: set[str] = set()
        for node in element.iter():
            if node.tag != qn("a:blip"):
                continue
            relation_id = node.get(qn("r:embed"))
            if not relation_id or relation_id in seen_rids:
                continue
            seen_rids.add(relation_id)
            part = doc.part.related_parts.get(relation_id)
            payload = getattr(part, "blob", None)
            if payload:
                payloads.append(payload)
        return payloads

    def _extract_markdown_image_sources(self, line: str) -> list[str]:
        return [match.strip() for match in re.findall(r"!\[[^\]]*\]\(([^)]+)\)", line) if match.strip()]

    def _resolve_image_payload(self, source_path: Path, image_source: str) -> bytes | None:
        cleaned_source = image_source.strip().strip("<>").strip()
        if not cleaned_source:
            return None
        if cleaned_source.startswith("data:"):
            return self._decode_data_url_image(cleaned_source)

        source_without_title = re.split(r"\s+", cleaned_source, maxsplit=1)[0]
        parsed = urlparse(source_without_title)
        if parsed.scheme in {"http", "https", "file"}:
            return None
        base_dir = source_path.parent.resolve()
        candidate = (base_dir / unquote(parsed.path or source_without_title)).resolve()
        try:
            candidate.relative_to(base_dir)
        except ValueError:
            return None
        try:
            if candidate.is_file():
                return candidate.read_bytes()
        except OSError:
            return None
        return None

    @staticmethod
    def _decode_data_url_image(source: str) -> bytes | None:
        match = re.match(r"^data:[^;]+;base64,(.+)$", source, flags=re.IGNORECASE | re.DOTALL)
        if not match:
            return None
        try:
            return base64.b64decode(match.group(1), validate=False)
        except Exception:
            return None

    def _filter_image_ocr_result(self, ocr_result: OcrResult) -> OcrResult:
        normalized_text = self._normalize_text(ocr_result.text)
        if not normalized_text:
            return replace(ocr_result, text="")
        if not self.settings.ocr_filter_noise_text:
            return replace(ocr_result, text=normalized_text)
        if self._should_keep_image_ocr_text(normalized_text, ocr_result):
            return replace(ocr_result, text=normalized_text)
        return replace(ocr_result, text="")

    def _should_keep_image_ocr_text(self, text: str, ocr_result: OcrResult) -> bool:
        signal_char_count = self._count_signal_characters(text)
        signal_unit_count = self._count_signal_units(text)
        if self._looks_like_noise_only_text(text):
            return False
        if signal_char_count < self.settings.ocr_image_min_text_chars and signal_unit_count < self.settings.ocr_image_min_tokens:
            return False
        if ocr_result.layout_hint == "low_signal" and not ocr_result.tables:
            return False
        return True

    @staticmethod
    def _count_signal_characters(text: str) -> int:
        return len(re.findall(r"[A-Za-z0-9\u4e00-\u9fff]", text))

    @staticmethod
    def _count_signal_units(text: str) -> int:
        ascii_tokens = re.findall(r"[A-Za-z0-9]+(?:[-_./][A-Za-z0-9]+)*", text)
        cjk_characters = re.findall(r"[\u4e00-\u9fff]", text)
        return len(ascii_tokens) + len(cjk_characters)

    @staticmethod
    def _looks_like_noise_only_text(text: str) -> bool:
        collapsed = re.sub(r"\s+", "", text)
        if not collapsed:
            return True
        noise_patterns = (
            r"^(图|表|附图|附件)\s*[:：-]?\s*[\dA-Za-z一二三四五六七八九十]+$",
            r"^(FIG(?:URE)?|IMAGE)\s*[:：-]?\s*[\dA-Za-z._-]+$",
            r"^(第?\d+页|\d+/\d+)$",
            r"^\d{4}[-/.年]\d{1,2}(?:[-/.月]\d{1,2}日?)?$",
            r"^\d+$",
        )
        return any(re.fullmatch(pattern, collapsed, flags=re.IGNORECASE) for pattern in noise_patterns)

    @staticmethod
    def _ocr_result_has_meaningful_content(ocr_result: OcrResult) -> bool:
        return bool(ocr_result.text.strip() or ocr_result.tables)

    @staticmethod
    def _next_embedded_image_index(segments: list[ParsedSegment], section_title_prefix: str) -> int:
        pattern = re.compile(rf"^{re.escape(section_title_prefix)} (\d+)(?:\s|$)")
        indexes = []
        for segment in segments:
            if not segment.section_title:
                continue
            match = pattern.match(segment.section_title)
            if match:
                indexes.append(int(match.group(1)))
        return (max(indexes) if indexes else 0) + 1

    def _extract_html_table_rows(self, table_element) -> list[list[str]]:
        rows: list[list[str]] = []
        for row in table_element.find_all("tr"):
            cells = row.find_all(["th", "td"], recursive=False)
            if not cells:
                cells = row.find_all(["th", "td"])
            rows.append([self._normalize_text(cell.get_text(" ", strip=True)) for cell in cells])
        return rows

    def _extract_pdf_table_segments(self, path: Path) -> dict[int, list[tuple[str, str]]]:
        extraction = self._extract_pdf_with_pdfplumber(path)
        return extraction.table_segments_by_page if extraction else {}

    def _extract_pdf_with_pdfplumber(self, path: Path) -> PdfPlumberExtraction | None:
        if pdfplumber is None:
            return None

        page_texts: dict[int, str] = {}
        segments_by_page: dict[int, list[tuple[str, str]]] = {}
        previous_headers_by_width: dict[int, tuple[list[str], int]] = {}
        try:
            with pdfplumber.open(str(path)) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    tables = self._extract_pdfplumber_tables(page)
                    table_bboxes = [bbox for _, bbox in tables if bbox is not None]
                    page_texts[page_number] = self._extract_pdfplumber_text_outside_tables(page, table_bboxes)
                    for table_index, (rows, _) in enumerate(tables, start=1):
                        if not rows:
                            continue
                        cleaned_rows = [
                            [self._normalize_pdf_text(cell or "") for cell in row]
                            for row in rows
                            if row
                        ]
                        cleaned_rows = self._prepare_pdf_table_rows(
                            cleaned_rows,
                            page_number=page_number,
                            previous_headers_by_width=previous_headers_by_width,
                        )
                        caption = f"PDF page {page_number} table {table_index}"
                        table_segments = table_rows_to_text_segments(cleaned_rows, caption=caption)
                        for table_segment in table_segments:
                            segments_by_page.setdefault(page_number, []).append((table_segment, caption))
        except Exception:
            return None

        return PdfPlumberExtraction(
            page_count=len(page_texts),
            page_texts=page_texts,
            table_segments_by_page=segments_by_page,
        )

    def _extract_pdfplumber_tables(self, page) -> list[tuple[list[list[str | None]], tuple[float, float, float, float] | None]]:
        tables: list[tuple[list[list[str | None]], tuple[float, float, float, float] | None]] = []
        if hasattr(page, "find_tables"):
            try:
                for table in page.find_tables() or []:
                    rows = table.extract() or []
                    bbox = getattr(table, "bbox", None)
                    tables.append((rows, bbox))
                if tables:
                    return tables
            except Exception:
                tables = []

        try:
            return [(rows, None) for rows in (page.extract_tables() or [])]
        except Exception:
            return []

    def _extract_pdfplumber_text_outside_tables(self, page, table_bboxes: list[tuple[float, float, float, float]]) -> str:
        if not hasattr(page, "extract_text"):
            return ""

        if table_bboxes and hasattr(page, "filter"):
            try:
                filtered_page = page.filter(lambda obj: not self._pdf_object_overlaps_any_bbox(obj, table_bboxes))
                return self._normalize_pdf_text(filtered_page.extract_text() or "")
            except Exception:
                return ""

        try:
            return self._normalize_pdf_text(page.extract_text() or "")
        except Exception:
            return ""

    def _pdf_object_overlaps_any_bbox(self, obj, bboxes: list[tuple[float, float, float, float]]) -> bool:
        obj_bbox = self._pdf_object_bbox(obj)
        if obj_bbox is None:
            return False
        return any(self._bboxes_overlap(obj_bbox, table_bbox) for table_bbox in bboxes)

    @staticmethod
    def _pdf_object_bbox(obj) -> tuple[float, float, float, float] | None:
        try:
            x0 = float(obj["x0"])
            x1 = float(obj["x1"])
            top = float(obj.get("top", obj.get("y0")))
            bottom = float(obj.get("bottom", obj.get("y1")))
        except (KeyError, TypeError, ValueError):
            return None
        return (x0, top, x1, bottom)

    @staticmethod
    def _bboxes_overlap(
        first: tuple[float, float, float, float],
        second: tuple[float, float, float, float],
    ) -> bool:
        first_x0, first_top, first_x1, first_bottom = first
        second_x0, second_top, second_x1, second_bottom = second
        return first_x0 < second_x1 and first_x1 > second_x0 and first_top < second_bottom and first_bottom > second_top

    def _prepare_pdf_table_rows(
        self,
        rows: list[list[str]],
        *,
        page_number: int,
        previous_headers_by_width: dict[int, tuple[list[str], int]],
    ) -> list[list[str]]:
        cleaned_rows = [row for row in rows if any(row)]
        if not cleaned_rows:
            return []

        width = max(len(row) for row in cleaned_rows)
        aligned_rows = [self._pad_row(row, width) for row in cleaned_rows]
        first_row = aligned_rows[0]
        next_row = aligned_rows[1] if len(aligned_rows) > 1 else None

        if self._looks_like_pdf_header_row(first_row, next_row):
            previous_headers_by_width[width] = (first_row, page_number)
            return aligned_rows

        previous = previous_headers_by_width.get(width)
        if previous and page_number - previous[1] <= 1:
            return [previous[0], *aligned_rows]

        return aligned_rows

    @staticmethod
    def _pad_row(row: list[str], width: int) -> list[str]:
        if len(row) >= width:
            return row[:width]
        return [*row, *([""] * (width - len(row)))]

    @staticmethod
    def _looks_like_pdf_header_row(row: list[str], next_row: list[str] | None = None) -> bool:
        cells = [cell for cell in row if cell]
        if len(cells) < 2:
            return False

        header_keywords = (
            "类型",
            "场景",
            "审批",
            "审批人",
            "职责",
            "责任",
            "时限",
            "周期",
            "材料",
            "备注",
            "等级",
            "条件",
            "方式",
            "对象",
            "要求",
            "字段",
            "问题",
            "回答",
            "依据",
            "是否",
            "负责人",
            "完成",
            "证据",
            "影响",
            "角色",
        )
        keyword_hits = sum(any(keyword in cell for keyword in header_keywords) for cell in cells)
        if keyword_hits >= max(1, len(cells) // 3):
            return True

        short_cells = sum(len(cell) <= 18 for cell in cells)
        data_like_cells = sum(bool(re.search(r"\d|万元|工作日|小时|天|年|L[1-4]|P[1-4]|是|否", cell)) for cell in cells)
        if next_row:
            next_data_like_cells = sum(bool(re.search(r"\d|万元|工作日|小时|天|年|L[1-4]|P[1-4]|是|否", cell)) for cell in next_row if cell)
            if short_cells >= max(2, len(cells) - 1) and next_data_like_cells > data_like_cells:
                return True

        return short_cells == len(cells) and data_like_cells == 0

    @staticmethod
    def _read_text_file(path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="ignore")

    @staticmethod
    def _split_plain_paragraphs(text: str) -> list[str]:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        raw_parts = re.split(r"\n\s*\n+", normalized)
        paragraphs = [DocumentParser._normalize_text(part) for part in raw_parts]
        return [paragraph for paragraph in paragraphs if paragraph]

    @staticmethod
    def _strip_markdown_inline(text: str) -> str:
        text = re.sub(r"!\[[^\]]*\]\([^)]*\)", " ", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
        text = re.sub(r"[`*_~>#]", "", text)
        text = re.sub(r"^\s*[-*+]\s+", "", text)
        text = re.sub(r"^\s*\d+\.\s+", "", text)
        return DocumentParser._normalize_text(text)

    @classmethod
    def _is_markdown_table_start(cls, lines: list[str], index: int) -> bool:
        if index + 1 >= len(lines):
            return False
        header = lines[index].strip()
        separator = lines[index + 1].strip()
        return "|" in header and cls._is_markdown_separator_row(separator)

    @classmethod
    def _collect_markdown_table(cls, lines: list[str], index: int) -> tuple[list[list[str]], int]:
        rows = [cls._parse_markdown_table_row(lines[index])]
        index += 2
        while index < len(lines):
            stripped = lines[index].strip()
            if not stripped or "|" not in stripped:
                break
            rows.append(cls._parse_markdown_table_row(stripped))
            index += 1
        return rows, index

    @staticmethod
    def _is_markdown_separator_row(line: str) -> bool:
        cells = DocumentParser._parse_markdown_table_row(line)
        if not cells:
            return False
        return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)

    @staticmethod
    def _parse_markdown_table_row(line: str) -> list[str]:
        stripped = line.strip().strip("|")
        return [DocumentParser._strip_markdown_inline(cell.strip()) for cell in stripped.split("|")]

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\x00", "")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @classmethod
    def _normalize_pdf_text(cls, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        raw_lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
        paragraphs: list[str] = []
        current = ""

        def flush_current() -> None:
            nonlocal current
            if current:
                paragraphs.append(current)
                current = ""

        for line in raw_lines:
            if not line:
                flush_current()
                continue
            if cls._looks_like_pdf_heading_line(line):
                flush_current()
                paragraphs.append(line)
                continue
            if not current:
                current = line
                continue
            separator = "" if cls._should_join_pdf_lines(current, line) else " "
            current = f"{current}{separator}{line}"

        flush_current()
        return "\n\n".join(DocumentParser._normalize_text(paragraph) for paragraph in paragraphs if paragraph)

    @staticmethod
    def _looks_like_pdf_heading_line(line: str) -> bool:
        return bool(
            re.match(r"^([一二三四五六七八九十百]+、|第[一二三四五六七八九十百0-9]+[章节])", line)
            or line in {"附则"}
        )

    @staticmethod
    def _should_join_pdf_lines(previous: str, current: str) -> bool:
        if not previous or not current:
            return False
        if previous.endswith("-"):
            return True
        if re.search(r"[\u4e00-\u9fff]$", previous) and re.match(r"^[\u4e00-\u9fff]", current):
            return not previous.endswith(("。", "；", "：", "！", "？"))
        return False

    def _finalize_segments(self, segments: list[ParsedSegment], parser_name: str, page_count: int | None = None) -> ParsedDocument:
        normalized_segments = [replace(segment, text=self._normalize_text(segment.text)) for segment in segments if self._normalize_text(segment.text)]
        combined_parts: list[str] = []
        cursor = 0
        finalized_segments: list[ParsedSegment] = []

        for segment in normalized_segments:
            if combined_parts:
                cursor += 2
            combined_parts.append(segment.text)
            finalized_segments.append(
                replace(
                    segment,
                    char_start=cursor,
                    char_end=cursor + len(segment.text),
                )
            )
            cursor += len(segment.text)

        normalized_text = "\n\n".join(combined_parts)
        return ParsedDocument(
            normalized_text=normalized_text,
            segments=finalized_segments,
            page_count=page_count,
            parser_name=parser_name,
        )
