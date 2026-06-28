from __future__ import annotations

import time
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.services.ingestion.markitdown_parser import MarkItDownParser
from app.services.ingestion import parsers as parser_module
from app.services.ingestion.parsers import DocumentParser


def _write_minimal_pdf(path: Path, text: str) -> None:
    stream_text = f"BT /F1 18 Tf 50 100 Td ({text}) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        f"<< /Length {len(stream_text)} >>\nstream\n{stream_text}\nendstream".encode("utf-8"),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += f"{index} 0 obj\n".encode("utf-8") + obj + b"\nendobj\n"
    xref_offset = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n".encode("utf-8")
    pdf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode("utf-8")
    pdf += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("utf-8")
    )
    path.write_bytes(pdf)


def test_document_parser_supports_multiple_formats(tmp_path: Path) -> None:
    parser = DocumentParser()

    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("Alpha paragraph.\n\nBeta paragraph.", encoding="utf-8")

    md_path = tmp_path / "notes.md"
    md_path.write_text("# Overview\n\nThis is a markdown paragraph.", encoding="utf-8")

    html_path = tmp_path / "page.html"
    html_path.write_text("<html><body><h1>Portal</h1><p>HTML paragraph.</p></body></html>", encoding="utf-8")

    docx_path = tmp_path / "report.docx"
    doc = DocxDocument()
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph("DOCX paragraph for parsing.")
    doc.save(docx_path)

    pdf_path = tmp_path / "sample.pdf"
    _write_minimal_pdf(pdf_path, "PDF text")

    txt_result = parser.parse(txt_path)
    md_result = parser.parse(md_path)
    html_result = parser.parse(html_path)
    docx_result = parser.parse(docx_path)
    pdf_result = parser.parse(pdf_path)

    assert txt_result.segments[0].paragraph_index == 1
    assert md_result.segments[0].section_title == "Overview"
    assert any(segment.section_title == "Portal" for segment in html_result.segments)
    assert any(segment.section_title == "Quarterly Report" for segment in docx_result.segments)
    assert pdf_result.page_count == 1
    assert pdf_result.normalized_text


def test_document_parser_extracts_markdown_tables(tmp_path: Path) -> None:
    parser = DocumentParser()
    md_path = tmp_path / "policy.md"
    md_path.write_text(
        "\n".join(
            [
                "# Leave Rules",
                "",
                "| Type | Approver | Notice |",
                "| --- | --- | --- |",
                "| Annual leave | Direct manager | 3 business days |",
                "| Sick leave | Team lead | Same day |",
            ]
        ),
        encoding="utf-8",
    )

    result = parser.parse(md_path)

    assert "Table row: Leave Rules. Type=Annual leave; Approver=Direct manager; Notice=3 business days." in result.normalized_text
    assert "Type=Sick leave; Approver=Team lead; Notice=Same day." in result.normalized_text


def test_document_parser_extracts_html_tables(tmp_path: Path) -> None:
    parser = DocumentParser()
    html_path = tmp_path / "policy.html"
    html_path.write_text(
        """
        <html><body>
          <h1>Export Rules</h1>
          <table>
            <caption>Retention</caption>
            <tr><th>Data Type</th><th>Retention</th></tr>
            <tr><td>Audit logs</td><td>180 days</td></tr>
          </table>
        </body></html>
        """,
        encoding="utf-8",
    )

    result = parser.parse(html_path)

    assert "Table row: Retention. Data Type=Audit logs; Retention=180 days." in result.normalized_text


def test_document_parser_prefers_main_html_content_and_skips_navigation(tmp_path: Path) -> None:
    parser = DocumentParser()
    html_path = tmp_path / "policy.html"
    html_path.write_text(
        """
        <html>
          <body>
            <nav><a>首页</a><a>登录</a></nav>
            <header><p>站点导航</p></header>
            <main>
              <h1>数据安全管理办法</h1>
              <p>重要数据处理活动应当建立风险评估和审批记录。</p>
            </main>
            <footer><p>版权所有</p></footer>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    result = parser.parse(html_path)

    assert "数据安全管理办法" in result.normalized_text
    assert "重要数据处理活动应当建立风险评估和审批记录" in result.normalized_text
    assert "首页" not in result.normalized_text
    assert "登录" not in result.normalized_text
    assert "版权所有" not in result.normalized_text


def test_document_parser_skips_policy_metadata_tables(tmp_path: Path) -> None:
    parser = DocumentParser()
    html_path = tmp_path / "policy.html"
    html_path.write_text(
        """
        <html><body>
          <div class="pages_content">
            <table>
              <tr><td>标 题：</td><td>生成式人工智能服务管理暂行办法</td></tr>
              <tr><td>发文机关：</td><td>国家网信办</td></tr>
              <tr><td>发文字号：</td><td>第15号</td></tr>
              <tr><td>主题分类：</td><td>科技</td></tr>
            </table>
            <p>提供者应当依法承担网络信息内容生产者责任，履行网络信息安全义务。</p>
          </div>
        </body></html>
        """,
        encoding="utf-8",
    )

    result = parser.parse(html_path)

    assert "提供者应当依法承担网络信息内容生产者责任" in result.normalized_text
    assert "Table row" not in result.normalized_text
    assert "发文字号" not in result.normalized_text


def test_document_parser_extracts_docx_tables_in_body_order(tmp_path: Path) -> None:
    parser = DocumentParser()
    docx_path = tmp_path / "runbook.docx"
    doc = DocxDocument()
    doc.add_heading("Incident Runbook", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Severity"
    table.cell(0, 1).text = "Response"
    table.cell(1, 0).text = "P1"
    table.cell(1, 1).text = "5 minutes"
    doc.add_paragraph("After table paragraph.")
    doc.save(docx_path)

    result = parser.parse(docx_path)

    normalized = result.normalized_text
    table_text = "Table row: Incident Runbook. Severity=P1; Response=5 minutes."
    assert table_text in normalized
    assert normalized.index("Incident Runbook") < normalized.index(table_text)
    assert normalized.index(table_text) < normalized.index("After table paragraph.")


def test_document_parser_extracts_pdf_tables(tmp_path: Path, monkeypatch) -> None:
    class FakePage:
        def extract_tables(self):
            return [
                [
                    ["Data Type", "Approver", "SLA"],
                    ["Customer phone", "Admin", "2 business days"],
                ]
            ]

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

    class FakePdfPlumber:
        @staticmethod
        def open(path: str):
            return FakePdf()

    monkeypatch.setattr(parser_module, "pdfplumber", FakePdfPlumber)

    parser = DocumentParser()
    pdf_path = tmp_path / "table.pdf"
    _write_minimal_pdf(pdf_path, "PDF table source")

    result = parser.parse(pdf_path)

    assert result.parser_name == "pdf"
    assert result.page_count == 1
    assert "Table row: PDF page 1 table 1. Data Type=Customer phone; Approver=Admin; SLA=2 business days." in result.normalized_text


def test_document_parser_reuses_previous_pdf_table_header_for_continuation_page(monkeypatch) -> None:
    class FakePage:
        def __init__(self, tables):
            self._tables = tables

        def extract_tables(self):
            return self._tables

    class FakePdf:
        pages = [
            FakePage(
                [
                    [
                        ["Data Type", "Approver", "SLA"],
                        ["Customer phone", "Admin", "2 business days"],
                    ]
                ]
            ),
            FakePage(
                [
                    [
                        ["Audit logs", "Security owner", "180 days"],
                    ]
                ]
            ),
        ]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

    class FakePdfPlumber:
        @staticmethod
        def open(path: str):
            return FakePdf()

    monkeypatch.setattr(parser_module, "pdfplumber", FakePdfPlumber)

    parser = DocumentParser()
    segments_by_page = parser._extract_pdf_table_segments(Path("ignored.pdf"))

    continuation_segment = segments_by_page[2][0][0]
    assert "Data Type=Audit logs" in continuation_segment
    assert "Approver=Security owner" in continuation_segment
    assert "SLA=180 days" in continuation_segment


def test_document_parser_excludes_pdf_table_area_from_plain_text(monkeypatch) -> None:
    class FakeTable:
        bbox = (0, 100, 500, 220)

        def extract(self):
            return [
                ["Data Type", "Approver", "SLA"],
                ["Customer phone", "Admin", "2 business days"],
            ]

    class FakeFilteredPage:
        def extract_text(self):
            return "Policy intro outside the table."

    class FakePage:
        def find_tables(self):
            return [FakeTable()]

        def filter(self, predicate):
            return FakeFilteredPage()

        def extract_text(self):
            return "Policy intro outside the table.\nData Type Approver SLA Customer phone Admin 2 business days"

    class FakePdf:
        pages = [FakePage()]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

    class FakePdfPlumber:
        @staticmethod
        def open(path: str):
            return FakePdf()

    monkeypatch.setattr(parser_module, "pdfplumber", FakePdfPlumber)

    result = DocumentParser().parse(Path("table.pdf"))

    assert "Policy intro outside the table." in result.normalized_text
    assert "Data Type Approver SLA Customer phone Admin" not in result.normalized_text
    assert "Table row: PDF page 1 table 1. Data Type=Customer phone; Approver=Admin; SLA=2 business days." in result.normalized_text


def test_document_parser_extracts_csv_tables(tmp_path: Path) -> None:
    parser = DocumentParser()
    csv_path = tmp_path / "approvals.csv"
    csv_path.write_text("Request,Approver,SLA\nData export,Admin,1 day\nRefund,Manager,2 days\n", encoding="utf-8")

    result = parser.parse(csv_path)

    assert result.parser_name == "csv"
    assert "Table row: approvals. Request=Data export; Approver=Admin; SLA=1 day." in result.normalized_text
    assert "Request=Refund; Approver=Manager; SLA=2 days." in result.normalized_text


def test_document_parser_routes_xlsx_to_markitdown_table_segments(tmp_path: Path, monkeypatch) -> None:
    def fake_convert(self, path: Path) -> str:
        return "\n".join(
            [
                "## Budget",
                "",
                "| Item | Owner |",
                "| --- | --- |",
                "| Cloud | Platform |",
            ]
        )

    monkeypatch.setattr(MarkItDownParser, "_convert_local_to_markdown", fake_convert)
    xlsx_path = tmp_path / "budget.xlsx"
    xlsx_path.write_bytes(b"fake xlsx payload")

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(xlsx_path)

    assert result.parser_name == "markitdown:xlsx"
    assert "Sheet: Budget" in result.normalized_text
    assert "Table row: Sheet: Budget. Item=Cloud; Owner=Platform." in result.normalized_text
    assert result.segments[0].section_title == "Sheet: Budget"


def test_document_parser_preserves_multiple_markitdown_xlsx_sheets(tmp_path: Path, monkeypatch) -> None:
    def fake_convert(self, path: Path) -> str:
        return "\n".join(
            [
                "## HR",
                "",
                "| Policy | Owner |",
                "| --- | --- |",
                "| Onboarding | People |",
                "",
                "## Finance",
                "",
                "| Policy | Owner |",
                "| --- | --- |",
                "| Expense | Accounting |",
            ]
        )

    monkeypatch.setattr(MarkItDownParser, "_convert_local_to_markdown", fake_convert)
    xlsx_path = tmp_path / "multi.xlsx"
    xlsx_path.write_bytes(b"fake xlsx payload")

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(xlsx_path)

    assert "Sheet: HR" in result.normalized_text
    assert "Sheet: Finance" in result.normalized_text
    assert "Table row: Sheet: HR. Policy=Onboarding; Owner=People." in result.normalized_text
    assert "Table row: Sheet: Finance. Policy=Expense; Owner=Accounting." in result.normalized_text


def test_document_parser_normalizes_real_markitdown_xlsx_merged_title_rows(tmp_path: Path) -> None:
    openpyxl = pytest.importorskip("openpyxl")

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Budget 2026"
    worksheet["A1"] = "Budget Summary"
    worksheet.merge_cells("A1:C1")
    worksheet.append(["Workstream", "Owner", "Spend"])
    worksheet.append(["Platform", "Mei", 125000])
    worksheet.append(["Security", "Jun", 98000])
    risk_sheet = workbook.create_sheet("Risks")
    risk_sheet.append(["Risk", "Severity", "Mitigation"])
    risk_sheet.append(["Vendor delay", "High", "Dual source"])

    xlsx_path = tmp_path / "complex.xlsx"
    workbook.save(xlsx_path)

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(xlsx_path)

    assert result.parser_name == "markitdown:xlsx"
    assert "Unnamed:" not in result.normalized_text
    assert (
        "Table row: Sheet: Budget 2026 / Budget Summary. Workstream=Platform; Owner=Mei; Spend=125000."
        in result.normalized_text
    )
    budget_row = next(segment for segment in result.segments if "Workstream=Platform" in segment.text)
    assert budget_row.section_title == "Sheet: Budget 2026 / Budget Summary"
    assert budget_row.citation_metadata["sheet_name"] == "Budget 2026"
    assert budget_row.citation_metadata["table_row_index"] == 1
    assert budget_row.citation_metadata["table_headers"] == ["Workstream", "Owner", "Spend"]


def test_document_parser_routes_xls_to_markitdown_best_effort(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(
        MarkItDownParser,
        "_convert_local_to_markdown",
        lambda self, path: "## Legacy\n\n| Field | Value |\n| --- | --- |\n| SLA | 2 days |",
    )
    xls_path = tmp_path / "legacy.xls"
    xls_path.write_bytes(b"fake xls payload")

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(xls_path)

    assert result.parser_name == "markitdown:xls"
    assert "Table row: Sheet: Legacy. Field=SLA; Value=2 days." in result.normalized_text


def test_document_parser_extracts_real_markitdown_xls_sheet_metadata(tmp_path: Path) -> None:
    xlwt = pytest.importorskip("xlwt")

    workbook = xlwt.Workbook()
    budget_sheet = workbook.add_sheet("Legacy Budget")
    budget_sheet.write_merge(0, 0, 0, 2, "Budget Summary")
    budget_sheet.write(1, 0, "Workstream")
    budget_sheet.write(1, 1, "Owner")
    budget_sheet.write(1, 2, "Spend")
    budget_sheet.write(2, 0, "Platform")
    budget_sheet.write(2, 1, "Mei")
    budget_sheet.write(2, 2, 125000)
    budget_sheet.write(3, 0, "Security")
    budget_sheet.write(3, 1, "Jun")
    budget_sheet.write(3, 2, 98000)
    risk_sheet = workbook.add_sheet("Risks")
    risk_sheet.write(0, 0, "Risk")
    risk_sheet.write(0, 1, "Severity")
    risk_sheet.write(0, 2, "Mitigation")
    risk_sheet.write(1, 0, "Vendor delay")
    risk_sheet.write(1, 1, "High")
    risk_sheet.write(1, 2, "Dual source")

    xls_path = tmp_path / "legacy-complex.xls"
    workbook.save(str(xls_path))

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(xls_path)

    assert result.parser_name == "markitdown:xls"
    assert "Table row: Sheet: Legacy Budget / Budget Summary. Workstream=Platform; Owner=Mei; Spend=125000." in result.normalized_text
    assert "Unnamed:" not in result.normalized_text
    budget_row = next(segment for segment in result.segments if "Workstream=Platform" in segment.text)
    assert budget_row.section_title == "Sheet: Legacy Budget / Budget Summary"
    assert budget_row.citation_metadata["sheet_name"] == "Legacy Budget"
    assert budget_row.citation_metadata["table_headers"] == ["Workstream", "Owner", "Spend"]


def test_document_parser_routes_pptx_to_markitdown_segments(tmp_path: Path, monkeypatch) -> None:
    def fake_convert(self, path: Path) -> str:
        return "\n".join(
            [
                "<!-- Slide number: 1 -->",
                "",
                "# Launch Plan",
                "",
                "Release window is Friday.",
                "",
                "| Workstream | Owner |",
                "| --- | --- |",
                "| QA | Mei |",
            ]
        )

    monkeypatch.setattr(MarkItDownParser, "_convert_local_to_markdown", fake_convert)
    pptx_path = tmp_path / "deck.pptx"
    pptx_path.write_bytes(b"fake pptx payload")

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(pptx_path)

    assert result.parser_name == "markitdown:pptx"
    assert "Slide 1: Launch Plan" in result.normalized_text
    assert "Release window is Friday." in result.normalized_text
    assert "Table row: Slide 1: Launch Plan. Workstream=QA; Owner=Mei." in result.normalized_text
    assert "Slide number" not in result.normalized_text
    assert result.page_count is None


def test_document_parser_extracts_real_markitdown_pptx_slide_metadata(tmp_path: Path) -> None:
    pptx = pytest.importorskip("pptx")
    presentation = pptx.Presentation()
    slide1 = presentation.slides.add_slide(presentation.slide_layouts[0])
    slide1.shapes.title.text = "Launch Plan"
    slide1.placeholders[1].text = "Q3 release\nRegional rollout"

    slide2 = presentation.slides.add_slide(presentation.slide_layouts[5])
    slide2.shapes.title.text = "Execution Matrix"
    textbox = slide2.shapes.add_textbox(pptx.util.Inches(0.8), pptx.util.Inches(1.5), pptx.util.Inches(4.5), pptx.util.Inches(1.4))
    text_frame = textbox.text_frame
    text_frame.text = "Milestones"
    text_frame.add_paragraph().text = "Alpha sign-off by July 10"
    text_frame.add_paragraph().text = "Pilot by August 1"
    table = slide2.shapes.add_table(3, 2, pptx.util.Inches(5.1), pptx.util.Inches(1.4), pptx.util.Inches(4.1), pptx.util.Inches(1.2)).table
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "QA"
    table.cell(1, 1).text = "On track"
    table.cell(2, 0).text = "Ops"
    table.cell(2, 1).text = "Blocked"

    pptx_path = tmp_path / "complex.pptx"
    presentation.save(pptx_path)

    result = DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(pptx_path)

    assert result.parser_name == "markitdown:pptx"
    assert "Slide 2: Execution Matrix" in result.normalized_text
    assert "Table row: Slide 2: Execution Matrix. Owner=QA; Status=On track." in result.normalized_text
    assert "Milestones" in result.normalized_text
    execution_row = next(segment for segment in result.segments if "Owner=QA; Status=On track." in segment.text)
    assert execution_row.citation_metadata["slide_number"] == 2
    assert execution_row.citation_metadata["slide_title"] == "Execution Matrix"
    assert execution_row.citation_metadata["segment_kind"] == "table_row"


def test_document_parser_markitdown_disabled_rejects_office_suffix(tmp_path: Path, monkeypatch) -> None:
    parser = DocumentParser(markitdown_allowed_base_dir=tmp_path)
    monkeypatch.setattr(parser.settings, "markitdown_enabled", False)
    xlsx_path = tmp_path / "disabled.xlsx"
    xlsx_path.write_bytes(b"fake xlsx payload")

    with pytest.raises(ValueError, match="Unsupported parser"):
        parser.parse(xlsx_path)


def test_document_parser_reports_markitdown_conversion_failure(tmp_path: Path, monkeypatch) -> None:
    def fail_convert(self, path: Path) -> str:
        raise ValueError("conversion failed")

    monkeypatch.setattr(MarkItDownParser, "_convert_local_to_markdown", fail_convert)
    xlsx_path = tmp_path / "broken.xlsx"
    xlsx_path.write_bytes(b"fake xlsx payload")

    with pytest.raises(ValueError, match="conversion failed"):
        DocumentParser(markitdown_allowed_base_dir=tmp_path).parse(xlsx_path)


def test_markitdown_parser_rejects_paths_outside_allowed_base(tmp_path: Path) -> None:
    allowed_dir = tmp_path / "allowed"
    outside_dir = tmp_path / "outside"
    allowed_dir.mkdir()
    outside_dir.mkdir()
    outside_path = outside_dir / "blocked.xlsx"
    outside_path.write_bytes(b"fake xlsx payload")

    parser = MarkItDownParser(allowed_base_dir=allowed_dir)

    with pytest.raises(ValueError, match="configured document data directory"):
        parser.parse(outside_path, suffix=".xlsx")


def test_markitdown_parser_rejects_oversized_local_file(tmp_path: Path, monkeypatch) -> None:
    parser = MarkItDownParser(allowed_base_dir=tmp_path)
    monkeypatch.setattr(parser.settings, "markitdown_max_file_size_bytes", 4)
    xlsx_path = tmp_path / "oversized.xlsx"
    xlsx_path.write_bytes(b"12345")

    with pytest.raises(ValueError, match="exceeding the configured limit"):
        parser.parse(xlsx_path, suffix=".xlsx")


def test_markitdown_parser_enforces_output_char_limit(tmp_path: Path, monkeypatch) -> None:
    parser = MarkItDownParser(allowed_base_dir=tmp_path)
    monkeypatch.setattr(parser.settings, "markitdown_max_output_chars", 16)
    monkeypatch.setattr(parser, "_run_markitdown_conversion", lambda path: "# Budget\n\nA very long payload")
    xlsx_path = tmp_path / "large-output.xlsx"
    xlsx_path.write_bytes(b"fake xlsx payload")

    with pytest.raises(ValueError, match="exceeding the configured limit"):
        parser.parse(xlsx_path, suffix=".xlsx")


def test_markitdown_parser_times_out_slow_conversion(tmp_path: Path, monkeypatch) -> None:
    parser = MarkItDownParser(allowed_base_dir=tmp_path)
    monkeypatch.setattr(parser.settings, "markitdown_timeout_seconds", 0.01)

    def slow_conversion(path: Path) -> str:
        time.sleep(0.05)
        return "# Budget"

    monkeypatch.setattr(parser, "_run_markitdown_conversion", slow_conversion)
    xlsx_path = tmp_path / "slow.xlsx"
    xlsx_path.write_bytes(b"fake xlsx payload")

    with pytest.raises(ValueError, match="timed out"):
        parser.parse(xlsx_path, suffix=".xlsx")
