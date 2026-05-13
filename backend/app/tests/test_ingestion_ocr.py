from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.core.security import hash_password
from app.models.enums import RoleName
from app.models.role import Role
from app.models.user import User
from app.services.ingestion import image_tables
from app.services.ingestion import parsers as parser_module
from app.services.ingestion.ocr import OcrResult
from app.services.ingestion.parsers import DocumentParser


class FakeOcrService:
    def __init__(
        self,
        *,
        enabled: bool = True,
        image_result: OcrResult | None = None,
        pdf_result: OcrResult | None = None,
        bytes_result: OcrResult | None = None,
    ):
        self.enabled = enabled
        self.image_result = image_result or OcrResult.empty()
        self.pdf_result = pdf_result or OcrResult.empty()
        self.bytes_result = bytes_result or OcrResult.empty()
        self.image_paths: list[Path] = []
        self.pdf_pages: list[int] = []
        self.byte_payloads: list[bytes] = []

    def extract_image(self, path: Path) -> OcrResult:
        self.image_paths.append(path)
        return self.image_result if self.enabled else OcrResult.empty()

    def extract_pdf_page(self, path: Path, page_number: int) -> OcrResult:
        self.pdf_pages.append(page_number)
        return self.pdf_result if self.enabled else OcrResult.empty()

    def extract_bytes(self, payload: bytes) -> OcrResult:
        self.byte_payloads.append(payload)
        return self.bytes_result if self.enabled else OcrResult.empty()


def _write_blank_pdf(path: Path) -> None:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << >> >>",
        b"<< /Length 0 >>\nstream\n\nendstream",
    ]
    pdf = b"%PDF-1.4\n"
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf += f"{index} 0 obj\n".encode("utf-8") + obj + b"\nendobj\n"
    xref_offset = len(pdf)
    pdf += f"xref\n0 {len(objects) + 1}\n".encode("utf-8")
    pdf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        pdf += f"{offset:010d} 00000 n \n".encode("utf-8")
    pdf += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("utf-8")
    )
    path.write_bytes(pdf)


def _install_fake_pdfplumber(monkeypatch, pages):
    class FakePdf:
        def __init__(self, fake_pages):
            self.pages = fake_pages

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

    class FakePdfPlumber:
        @staticmethod
        def open(path: str):
            return FakePdf(pages)

    monkeypatch.setattr(parser_module, "pdfplumber", FakePdfPlumber)


SAMPLE_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO2Wk9sAAAAASUVORK5CYII="
)


def test_image_ocr_parser_returns_text_segment(tmp_path: Path) -> None:
    image_path = tmp_path / "notice.png"
    image_path.write_bytes(b"fake image bytes")
    ocr = FakeOcrService(image_result=OcrResult(text="扫描图片中的客户编号 ABC-123"))

    result = DocumentParser(ocr_service=ocr).parse(image_path)

    assert result.parser_name == "image"
    assert result.page_count == 1
    assert result.segments[0].section_title == "Image OCR"
    assert "客户编号 ABC-123" in result.normalized_text


def test_image_table_ocr_becomes_table_row_segment(tmp_path: Path) -> None:
    image_path = tmp_path / "table.jpg"
    image_path.write_bytes(b"fake image bytes")
    ocr = FakeOcrService(
        image_result=OcrResult(
            text="原始 OCR 文本",
            tables=[
                [
                    ["字段A", "字段B"],
                    ["值A", "值B"],
                ]
            ],
        )
    )

    result = DocumentParser(ocr_service=ocr).parse(image_path)

    assert "Table row: Image table 1. 字段A=值A; 字段B=值B." in result.normalized_text
    assert any(segment.section_title == "Image table 1" for segment in result.segments)


def test_low_signal_image_text_is_filtered_out(tmp_path: Path) -> None:
    image_path = tmp_path / "stamp.png"
    image_path.write_bytes(b"fake image bytes")
    parser = DocumentParser(
        ocr_service=FakeOcrService(image_result=OcrResult(text="图1 2026", layout_hint="low_signal"))
    )

    result = parser.parse(image_path)

    assert result.parser_name == "image"
    assert result.normalized_text == ""
    assert result.segments == []


def test_diagram_like_image_text_is_kept(tmp_path: Path) -> None:
    image_path = tmp_path / "flow.png"
    image_path.write_bytes(b"fake image bytes")
    parser = DocumentParser(
        ocr_service=FakeOcrService(
            image_result=OcrResult(
                text="审批发起 风控复核 系统授权 操作审计",
                layout_hint="diagram_like",
            )
        )
    )

    result = parser.parse(image_path)

    assert "审批发起" in result.normalized_text
    assert any(segment.section_title == "Image OCR" for segment in result.segments)


def test_low_signal_image_text_does_not_drop_detected_table(tmp_path: Path) -> None:
    image_path = tmp_path / "table_noise.png"
    image_path.write_bytes(b"fake image bytes")
    parser = DocumentParser(
        ocr_service=FakeOcrService(
            image_result=OcrResult(
                text="图1",
                tables=[
                    [
                        ["字段A", "字段B"],
                        ["值A", "值B"],
                    ]
                ],
                layout_hint="table_like",
            )
        )
    )

    result = parser.parse(image_path)

    assert "图1" not in result.normalized_text
    assert any("字段A=值A" in segment.text for segment in result.segments)


def test_markdown_embedded_local_image_is_ocr_processed(tmp_path: Path) -> None:
    image_path = tmp_path / "inline.png"
    image_path.write_bytes(SAMPLE_PNG_BYTES)
    markdown_path = tmp_path / "policy.md"
    markdown_path.write_text("# 说明\n\n![扫描件](inline.png)\n", encoding="utf-8")
    ocr = FakeOcrService(bytes_result=OcrResult(text="Markdown 图片 OCR 内容"))

    result = DocumentParser(ocr_service=ocr).parse(markdown_path)

    assert "Markdown 图片 OCR 内容" in result.normalized_text
    assert any(segment.section_title == "Markdown image 1 OCR" for segment in result.segments)
    assert len(ocr.byte_payloads) == 1


def test_html_embedded_data_image_is_ocr_processed(tmp_path: Path) -> None:
    html_path = tmp_path / "policy.html"
    data_url = "data:image/png;base64," + base64.b64encode(SAMPLE_PNG_BYTES).decode("ascii")
    html_path.write_text(f"<html><body><h1>制度</h1><img src=\"{data_url}\" /></body></html>", encoding="utf-8")
    ocr = FakeOcrService(bytes_result=OcrResult(text="HTML 图片 OCR 内容"))

    result = DocumentParser(ocr_service=ocr).parse(html_path)

    assert "HTML 图片 OCR 内容" in result.normalized_text
    assert any(segment.section_title == "HTML image 1 OCR" for segment in result.segments)
    assert len(ocr.byte_payloads) == 1


def test_docx_embedded_image_is_ocr_processed(tmp_path: Path) -> None:
    docx_path = tmp_path / "policy.docx"
    doc = DocxDocument()
    doc.add_heading("附件", level=1)
    doc.add_picture(BytesIO(SAMPLE_PNG_BYTES))
    doc.save(docx_path)
    ocr = FakeOcrService(bytes_result=OcrResult(text="DOCX 图片 OCR 内容"))

    result = DocumentParser(ocr_service=ocr).parse(docx_path)

    assert "DOCX 图片 OCR 内容" in result.normalized_text
    assert any(segment.section_title == "DOCX image 1 OCR" for segment in result.segments)
    assert len(ocr.byte_payloads) == 1


def test_image_table_extractor_groups_tesseract_data_into_rows(monkeypatch) -> None:
    class FakeOutput:
        DICT = "dict"

    class FakeTesseract:
        Output = FakeOutput

        @staticmethod
        def image_to_data(image, *, lang: str, output_type):
            return {
                "text": ["Name", "Amount", "Alice", "42"],
                "left": [10, 120, 10, 120],
                "top": [10, 10, 40, 40],
                "width": [40, 60, 40, 20],
                "height": [12, 12, 12, 12],
                "conf": ["96", "95", "94", "93"],
            }

    monkeypatch.setattr(image_tables, "pytesseract", FakeTesseract)

    rows = image_tables.extract_image_table_rows(object(), lang="eng")

    assert rows == [["Name", "Amount"], ["Alice", "42"]]


def test_image_table_extractor_prefers_contiguous_table_over_header_pairs(monkeypatch) -> None:
    class FakeOutput:
        DICT = "dict"

    class FakeTesseract:
        Output = FakeOutput

        @staticmethod
        def image_to_data(image, *, lang: str, output_type):
            return {
                "text": [
                    "申请编号",
                    "ACC-2026-0513-07",
                    "关联工单",
                    "CS-45821",
                    "编号",
                    "负责人",
                    "时限",
                    "动作",
                    "扫描A",
                    "张三",
                    "6小时",
                    "核验日志",
                    "扫描B",
                    "李四",
                    "12小时",
                    "关闭权限",
                ],
                "left": [10, 160, 10, 160, 10, 110, 210, 310, 10, 110, 210, 310, 10, 110, 210, 310],
                "top": [10, 10, 35, 35, 80, 80, 80, 80, 110, 110, 110, 110, 140, 140, 140, 140],
                "width": [60, 120, 60, 80, 40, 45, 35, 40, 50, 35, 40, 60, 50, 35, 50, 60],
                "height": [12] * 16,
                "conf": ["96"] * 16,
            }

    monkeypatch.setattr(image_tables, "pytesseract", FakeTesseract)

    rows = image_tables.extract_image_table_rows(object(), lang="chi_sim+eng")

    assert rows == [
        ["编号", "负责人", "时限", "动作"],
        ["扫描A", "张三", "6小时", "核验日志"],
        ["扫描B", "李四", "12小时", "关闭权限"],
    ]


def test_image_table_extractor_rejects_scattered_chart_labels(monkeypatch) -> None:
    class FakeOutput:
        DICT = "dict"

    class FakeTesseract:
        Output = FakeOutput

        @staticmethod
        def image_to_data(image, *, lang: str, output_type):
            return {
                "text": ["Q1", "12%", "Q2", "18%", "Q3", "26%", "Q4", "44%"],
                "left": [10, 210, 70, 280, 140, 210, 230, 320],
                "top": [10, 10, 50, 50, 95, 95, 145, 145],
                "width": [30, 36, 30, 36, 30, 36, 30, 36],
                "height": [16, 16, 16, 16, 16, 16, 16, 16],
                "conf": ["95", "95", "95", "95", "95", "95", "95", "95"],
            }

    monkeypatch.setattr(image_tables, "pytesseract", FakeTesseract)

    rows = image_tables.extract_image_table_rows(object(), lang="chi_sim")

    assert rows == []


def test_scanned_pdf_page_triggers_ocr_when_plain_text_is_empty(tmp_path: Path, monkeypatch) -> None:
    class EmptyPage:
        def extract_text(self):
            return ""

        def extract_tables(self):
            return []

    _install_fake_pdfplumber(monkeypatch, [EmptyPage()])
    pdf_path = tmp_path / "scan.pdf"
    _write_blank_pdf(pdf_path)
    ocr = FakeOcrService(
        pdf_result=OcrResult(
            text="扫描 PDF OCR 文本",
            tables=[
                [
                    ["项目", "金额"],
                    ["服务费", "42"],
                ]
            ],
        )
    )

    result = DocumentParser(ocr_service=ocr).parse(pdf_path)

    assert ocr.pdf_pages == [1]
    assert "扫描 PDF OCR 文本" in result.normalized_text
    assert "Table row: PDF page 1 OCR table 1. 项目=服务费; 金额=42." in result.normalized_text
    assert result.segments[0].page_number == 1


def test_text_pdf_page_still_ocr_processes_embedded_images(tmp_path: Path, monkeypatch) -> None:
    class TextPage:
        def extract_text(self):
            return "This page has plenty of embedded selectable text already."

        def extract_tables(self):
            return []

    _install_fake_pdfplumber(monkeypatch, [TextPage()])
    monkeypatch.setattr(DocumentParser, "_open_pdf_image_document", lambda self, path: object())
    monkeypatch.setattr(DocumentParser, "_extract_pdf_page_image_payloads", lambda self, doc, page_number: [SAMPLE_PNG_BYTES])
    pdf_path = tmp_path / "mixed.pdf"
    _write_blank_pdf(pdf_path)
    ocr = FakeOcrService(bytes_result=OcrResult(text="PDF 内嵌图片 OCR 文本"))

    result = DocumentParser(ocr_service=ocr).parse(pdf_path)

    assert ocr.pdf_pages == []
    assert "embedded selectable text already" in result.normalized_text
    assert "PDF 内嵌图片 OCR 文本" in result.normalized_text
    assert any(segment.section_title == "PDF page 1 image 1 OCR" for segment in result.segments)


def test_text_pdf_page_is_not_ocr_duplicated(tmp_path: Path, monkeypatch) -> None:
    class TextPage:
        def extract_text(self):
            return "This PDF page already has enough embedded selectable text for normal parsing."

        def extract_tables(self):
            return []

    _install_fake_pdfplumber(monkeypatch, [TextPage()])
    pdf_path = tmp_path / "text.pdf"
    _write_blank_pdf(pdf_path)
    ocr = FakeOcrService(pdf_result=OcrResult(text="OCR should not run"))

    result = DocumentParser(ocr_service=ocr).parse(pdf_path)

    assert ocr.pdf_pages == []
    assert "embedded selectable text" in result.normalized_text
    assert "OCR should not run" not in result.normalized_text


def test_ocr_disabled_returns_empty_result_without_crashing(tmp_path: Path, monkeypatch) -> None:
    image_path = tmp_path / "disabled.jpeg"
    image_path.write_bytes(b"fake image bytes")
    disabled_ocr = FakeOcrService(enabled=False)

    image_result = DocumentParser(ocr_service=disabled_ocr).parse(image_path)

    assert image_result.parser_name == "image"
    assert image_result.normalized_text == ""

    class EmptyPage:
        def extract_text(self):
            return ""

        def extract_tables(self):
            return []

    _install_fake_pdfplumber(monkeypatch, [EmptyPage()])
    pdf_path = tmp_path / "disabled-scan.pdf"
    _write_blank_pdf(pdf_path)

    pdf_result = DocumentParser(ocr_service=disabled_ocr).parse(pdf_path)

    assert pdf_result.parser_name == "pdf"
    assert pdf_result.page_count == 1
    assert pdf_result.normalized_text == ""
    assert disabled_ocr.pdf_pages == []


def _create_user(db_session: Session, role: Role, email: str, password: str) -> User:
    user = User(
        email=email,
        full_name=email.split("@")[0],
        password_hash=hash_password(password),
        is_active=True,
        role_id=role.id,
    )
    db_session.add(user)
    db_session.flush()
    return user


def _login(client: TestClient, email: str, password: str) -> str:
    response = client.post("/api/v1/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    return response.json()["access_token"]


def test_png_jpg_jpeg_uploads_are_allowed(client: TestClient, db_session: Session) -> None:
    admin_role = Role(name=RoleName.ADMIN, description="Admin")
    db_session.add(admin_role)
    db_session.flush()
    _create_user(db_session, admin_role, "admin@example.com", "admin-pass")
    db_session.commit()
    admin_token = _login(client, "admin@example.com", "admin-pass")

    for filename, mime_type in [
        ("scan.png", "image/png"),
        ("photo.jpg", "image/jpeg"),
        ("photo.jpeg", "image/jpeg"),
    ]:
        upload_response = client.post(
            "/api/v1/documents/upload",
            headers={"Authorization": f"Bearer {admin_token}"},
            files={"file": (filename, BytesIO(b"not a real image"), mime_type)},
            data={"title": filename, "description": "image upload", "status": "active"},
        )

        assert upload_response.status_code == 200
        assert upload_response.json()["version"]["original_filename"] == filename
